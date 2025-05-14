import tkinter as tk
from tkinter import ttk, messagebox
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_section_columns(section, num_columns=2, space_between=Pt(30)):
    """
    Define o número de colunas em uma seção do documento Word.
    """
    sectPr = section._sectPr
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), str(num_columns))
    cols.set(qn('w:space'), str(int(space_between.twips)))  # Espaço entre colunas
    sectPr.append(cols)

def gerar_documento(dados):
    try:
        # Criar um novo documento Word
        doc = Document()

        # Adicionar título
        title = doc.add_heading("EQUIPAMENTO EM MANUTENÇÃO", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.size = Pt(24)

        # Adicionar espaço em branco (2 parágrafos vazios)
        for _ in range(2):
            doc.add_paragraph()

        # Criar uma nova seção com duas colunas, sem quebra de página
        section = doc.add_section(start_type=WD_SECTION_START.CONTINUOUS)
        set_section_columns(section, num_columns=2, space_between=Pt(20))  # 20 pontos de espaço entre colunas

        # Adicionar campos da esquerda (Entrada até Observação)
        doc.add_paragraph()
        campos_esquerda = [
            ("Entrada", dados["entrada"]),
            ("Modelo", dados["modelo"]),
            ("Serial", dados["serial"]),
            ("Empresa", dados["empresa"]),
            ("Telefone", dados["telefone"]),
            ("Contato", dados["contato"]),
            ("ID", dados["id"]),
        ]

        # Calcular o tamanho máximo do nome do campo para alinhamento
        max_label_length = max(len(nome) for nome, _ in campos_esquerda + [("Defeito", ""), ("Observação", "")])
        for nome, valor in campos_esquerda:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(f"{nome}:")
            run.bold = True
            run.font.size = Pt(14)
            padding = " " * (max_label_length - len(nome) + 4)
            run = paragraph.add_run(f"{padding}{valor}")
            run.font.size = Pt(14)

        # Adicionar Defeito com valor na linha de baixo
        paragraph = doc.add_paragraph()
        run = paragraph.add_run("Defeito:")
        run.bold = True
        run.font.size = Pt(14)
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(dados["defeito"])
        run.font.size = Pt(14)

        # Adicionar Observação com valor na linha de baixo
        paragraph = doc.add_paragraph()
        run = paragraph.add_run("Observação:")
        run.bold = True
        run.font.size = Pt(14)
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(dados["obs"])
        run.font.size = Pt(14)

        # Adicionar uma quebra de coluna para forçar o próximo conteúdo na coluna da direita
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_break(WD_BREAK.COLUMN)

        # Adicionar campos da direita (Caixa de transporte, Garantia, Checklist)
        # Adicionar Caixa de transporte com checkboxes
        caixa_checked = dados["caixa"] == "Sim"
        paragraph = doc.add_paragraph()
        run = paragraph.add_run("Caixa de transporte:")
        run.bold = True
        run.font.size = Pt(14)
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(f"({'x' if caixa_checked else ' '}) Sim   ({'x' if not caixa_checked else ' '}) Não")
        run.font.size = Pt(14)

        # Adicionar Garantia com checkboxes
        sim_checked = dados["garantia"] == "Sim"
        paragraph = doc.add_paragraph()
        run = paragraph.add_run("Garantia:")
        run.bold = True
        run.font.size = Pt(14)
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(f"({'x' if sim_checked else ' '}) Sim   ({'x' if not sim_checked else ' '}) Não")
        run.font.size = Pt(14)

        # Adicionar linha em branco antes do checklist
        doc.add_paragraph()

        # Adicionar checklist
        checklist = [
            ("Inserido no sistema", dados["check1"]),
            ("Informado orçamento", dados["check2"]),
            ("Manutenção Aprovada", dados["check3"]),
            ("Aguardando peças", dados["check4"]),
            ("Pronto para entrega", dados["check5"]),
        ]

        paragraph = doc.add_paragraph()
        run = paragraph.add_run("Checklist:")
        run.bold = True
        run.font.size = Pt(14)

        for nome, checked in checklist:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(f"     {nome}:")
            run.font.size = Pt(14)  # Sem negrito
            padding = " " * (max_label_length - len(nome) + 8)
            run = paragraph.add_run(f"{padding}({'x' if checked else ' '})")
            run.font.size = Pt(14)

        # Salvar documento
        novo_arquivo = f"diagnostico_{dados['serial']}.docx"
        doc.save(novo_arquivo)
        messagebox.showinfo("Sucesso", f"Documento salvo como {novo_arquivo}!")

    except Exception as e:
        messagebox.showerror("Erro", f"Erro ao criar o documento: {e}")
        print(f"Erro detalhado: {e}")

def criar_interface():
    root = tk.Tk()
    root.title("Sistema de Diagnóstico de Equipamento")

    campos = {
        "entrada": tk.StringVar(),
        "modelo": tk.StringVar(),
        "serial": tk.StringVar(),
        "empresa": tk.StringVar(),
        "telefone": tk.StringVar(),
        "contato": tk.StringVar(),
        "id": tk.StringVar(),
        "caixa": tk.StringVar(value="Não"),
        "garantia": tk.StringVar(value="Não"),
        "defeito": tk.StringVar(),
        "obs": tk.StringVar(),
        "check1": tk.BooleanVar(),
        "check2": tk.BooleanVar(),
        "check3": tk.BooleanVar(),
        "check4": tk.BooleanVar(),
        "check5": tk.BooleanVar()
    }

    def enviar():
        dados = {k: v.get() if isinstance(v, (tk.StringVar, tk.BooleanVar)) else v for k, v in campos.items()}
        gerar_documento(dados)

    labels = [
        ("Data de Entrada", "entrada"),
        ("Modelo", "modelo"),
        ("Serial", "serial"),
        ("Empresa", "empresa"),
        ("Telefone", "telefone"),
        ("Contato", "contato"),
        ("ID", "id"),
    ]

    for idx, (label, key) in enumerate(labels):
        ttk.Label(root, text=label).grid(row=idx, column=0, sticky="w", padx=5, pady=2)
        ttk.Entry(root, textvariable=campos[key], width=50).grid(row=idx, column=1, pady=2)

    ttk.Label(root, text="Caixa de transporte").grid(row=7, column=0, sticky="w", padx=5, pady=2)
    frame_caixa = ttk.Frame(root)
    frame_caixa.grid(row=7, column=1, sticky="w")
    ttk.Radiobutton(frame_caixa, text="Sim", variable=campos["caixa"], value="Sim").pack(side="left")
    ttk.Radiobutton(frame_caixa, text="Não", variable=campos["caixa"], value="Não").pack(side="left")

    ttk.Label(root, text="Garantia").grid(row=8, column=0, sticky="w", padx=5, pady=2)
    frame_garantia = ttk.Frame(root)
    frame_garantia.grid(row=8, column=1, sticky="w")
    ttk.Radiobutton(frame_garantia, text="Sim", variable=campos["garantia"], value="Sim").pack(side="left")
    ttk.Radiobutton(frame_garantia, text="Não", variable=campos["garantia"], value="Não").pack(side="left")

    ttk.Label(root, text="Defeito").grid(row=9, column=0, sticky="w", padx=5, pady=2)
    ttk.Entry(root, textvariable=campos["defeito"], width=50).grid(row=9, column=1, pady=2)

    ttk.Label(root, text="Observação").grid(row=10, column=0, sticky="w", padx=5, pady=2)
    ttk.Entry(root, textvariable=campos["obs"], width=50).grid(row=10, column=1, pady=2)

    checklist = [
        ("Inserido no sistema", "check1"),
        ("Informado orçamento", "check2"),
        ("Manutenção Aprovada", "check3"),
        ("Aguardando peças", "check4"),
        ("Pronto para entrega", "check5")
    ]

    ttk.Label(root, text="Checklist:").grid(row=11, column=0, sticky="w", padx=5, pady=5)
    for i, (texto, chave) in enumerate(checklist):
        ttk.Checkbutton(root, text=texto, variable=campos[chave]).grid(row=11 + i, column=1, sticky="w")

    ttk.Button(root, text="Gerar Word", command=enviar).grid(row=16, column=1, pady=10)

    root.mainloop()

if __name__ == "__main__":
    criar_interface()